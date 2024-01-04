from pathlib import Path
from csv import DictReader
from docx import Document
from docx.shared import Inches, Pt, Mm, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from utils import read_csv, subnets, shade_cells

BASE_DIR = Path(__file__).resolve().parent
DATA_FOLDER = BASE_DIR / "data"
# CSV name (csv file should be in data folder - CVETG/data/)
CSV = DATA_FOLDER / "combined_hta_2023.csv"  # CSV file name


data = list(x for x in read_csv(CSV) if x['Host'].split('.')[2] in ['40','50','120','121']) # filter main data

# Create risk counts (critical, high, medium only)
risk_counts = [
    {
        "risk": "Critical",
        "count": len([x for x in data if x.get("Risk") == "Critical"]),
    },
    {
        "risk": "High",
        "count": len([x for x in data if x.get("Risk") == "High"]),
    },
    {
        "risk": "Medium",
        "count": len([x for x in data if x.get("Risk") == "Medium"]),
    }
]

# List of hosts (distinct ip addresses)
hosts = set([x.get("Host") for x in data])

# Create a new Document
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.bold = False
doc.styles['Normal'].paragraph_format.line_spacing = 1.5
section = doc.sections[0]
section.top_margin = Cm(2.25)
section.bottom_margin = Cm(2.25)
section.left_margin = Cm(1.9)
section.right_margin = Cm(2.25)
section.page_height = Mm(297)
section.page_width = Mm(210)


# Information table
table = doc.add_table(rows=5, cols=4)
table.style = "Light Grid Accent 1"
row1col1 = table.cell(0, 0).merge(table.cell(0, 1))
row1col1.paragraphs[0].add_run("Target Information").bold = True
table.cell(0, 2).paragraphs[0].add_run("Total Risk").bold = True
table.cell(0, 3).paragraphs[0].add_run(
    f"{max(risk_counts, key=lambda x: x['count']).get('risk')}"
).bold = True

table.cell(1, 0).paragraphs[0].add_run("Name").bold = True
table.cell(1, 2).merge(table.cell(1, 3)).paragraphs[0].add_run(
    "Engagement Information"
).bold = True
table.cell(2, 0).paragraphs[0].add_run("Type").bold = True
table.cell(2, 1).paragraphs[0].add_run(
    "Network Security Assessment"  # Assessment type name
).bold = False
table.cell(2, 2).paragraphs[0].add_run(
    "Security Auditor"
).bold = True
table.cell(2, 3).paragraphs[0].add_run("3").bold = False

table.cell(3, 0).paragraphs[0].add_run("Date").bold = True
table.cell(3, 1).paragraphs[0].add_run("").bold = False  # Date
table.cell(3, 2).paragraphs[0].add_run("Date").bold = True
table.cell(3, 3).paragraphs[0].add_run("30 Days").bold = False

table.cell(4, 0).merge(table.cell(4, 1)).paragraphs[0].add_run(
    f"Finding on {len(hosts)} IP addresses."
)
row4col2 = table.cell(4, 2).merge(table.cell(4, 3))
row4col2.paragraphs[0].add_run(
    "Subnets involve in assessment."
).bold = True
row4col2.add_paragraph().add_run(f"{subnets(hosts)}").bold = False

# Empty line
doc.add_paragraph()

VULNERABILITY_COUNT = 1
for row in data:
    if row.get("Risk") != "None":  # Only add rows with risk level of Critical, High, Medium or Low
        # Vulnerability tables
        table = doc.add_table(rows=8, cols=4)
        table.style = "Table Grid"
        table.allow_autofit = False

        # Row 1
        row_1_merged_cells = table.cell(0, 0).merge(table.cell(0, 1))
        row_1_merged_cells.paragraphs[0].add_run(
            f"{VULNERABILITY_COUNT}. {row.get('Name')}"
        ).bold = True
        table.cell(0, 2).paragraphs[0].add_run("IP Address").bold = True
        table.cell(0, 3).paragraphs[0].add_run(row.get("Host")).bold = True
        # Shade cells
        if row.get("Risk") == "Critical":
            shade_cells(cells=[table.cell(0, 3)], shade="DarkRed")
        elif row.get("Risk") == "High":
            shade_cells(cells=[table.cell(0, 3)], shade="Red")
        elif row.get("Risk") == "Medium":
            shade_cells(cells=[table.cell(0, 3)], shade="Orange")
        elif row.get("Risk") == "Low":
            shade_cells(cells=[table.cell(0, 3)], shade="#00B0F0")

        # Row 2
        row_2_merged_cells_1 = table.cell(1, 0).merge(table.cell(1, 1))
        row_2_merged_cells_2 = table.cell(1, 2).merge(table.cell(1, 3))
        row_2_merged_cells_1.paragraphs[0].add_run(
            "Type Of Pentest"
        ).bold = True
        row_2_merged_cells_2.paragraphs[0].add_run(
            "Network Security Assessments"  # Assessment type name
        ).bold = True

        # Row 3
        row_3_merged_cells_1 = table.cell(2, 0).merge(table.cell(2, 3))
        row_3_merged_cells_1.paragraphs[0].add_run("Description:").bold = True
        row_3_merged_cells_1.add_paragraph(row.get("Description"))
        row_3_merged_cells_1.add_paragraph()
        row_3_merged_cells_1.paragraphs[2].add_run("Protocol: ").bold = True
        row_3_merged_cells_1.paragraphs[2].add_run(
            f"{row.get('Protocol').upper()}"
        )
        row_3_merged_cells_1.add_paragraph()
        row_3_merged_cells_1.paragraphs[3].add_run("Port: ").bold = True
        row_3_merged_cells_1.paragraphs[3].add_run(
            f"{row.get('Port').upper()}"
        )

        # Row 4
        row_4_merged_cells_1 = table.cell(3, 0).merge(table.cell(3, 3))
        row_4_merged_cells_1.paragraphs[0].add_run(
            "Proof Of Concept (Scanner Output):").bold = True
        row_4_merged_cells_1.add_paragraph(row.get("Synopsis"))

        # Row 5
        row_5_merged_cells_1 = table.cell(4, 0).merge(table.cell(4, 3))
        row_5_merged_cells_1.paragraphs[0].add_run(
            "Recommendation:"
        ).bold = True
        row_5_merged_cells_1.add_paragraph(row.get("Solution"))

        # Row 6
        row_6_merged_cells_1 = table.cell(5, 0).merge(table.cell(5, 3))
        row_6_merged_cells_1.paragraphs[0].add_run("References:").bold = True
        row_6_merged_cells_1.add_paragraph(row.get("See Also"))

        # Row 7
        row_7_merged_cells_1 = table.cell(6, 0).merge(table.cell(6, 3))
        row_7_merged_cells_1.paragraphs[0].add_run("Action Taken:").bold = True
        row_7_merged_cells_1.add_paragraph().add_run("Date:").bold = True
        row_7_merged_cells_1.add_paragraph().add_run("By:").bold = True
        row_7_merged_cells_1.add_paragraph().add_run("Remarks:").bold = True

        # Row 8
        row_8_merged_cells_1 = table.cell(7, 0).merge(table.cell(7, 3))
        row_8_merged_cells_1.paragraphs[0].add_run("Status:").bold = True
        shade_cells([table.cell(7, 0)], "#92D050")

        # Empty line
        doc.add_paragraph()
        VULNERABILITY_COUNT += 1

# Save the document
doc.save(str(DATA_FOLDER/'Network Security Assessment.docx'))  # Output file name
